"""Parse a founder-filled ARTPARK MIS `.docx` (monthly update or quarterly
review) into `mis_catalog`'s own shape.

Builds on `template_parser.py`'s primitives conceptually (anchor-driven
extraction from a python-docx tree) but a fresh module: the MIS templates
have no `>>> ANSWER QN START >>>` markers at all — the real structure is
Heading 1/2 paragraphs, "List Paragraph"-styled bullets carrying a
`"<Label>: <value>"` prose line, and a handful of genuine Word tables
(`docs/reference/mis-templates.md` §1/§2, confirmed against the two real
source files under `tests/fixtures/mis_{monthly,quarterly}_template.docx`).

This module is pure — no DB, no storage, no FastAPI. `founder_mis.py`'s
import endpoints own I/O; this module only turns bytes into a
`ParsedMisDocument` plus a per-field record of what was matched (including a
real `None` for a field the document left genuinely blank), what was
ambiguous (present but could not be validated with confidence), and what was
ignored (document structure that maps to nothing in the catalog, or a
computed column read only to cross-check).

Three design rulings this module exists to enforce (constraints from the
build brief):

1. **Never invent a catalog key.** Every metric/financial-series/headcount
   row is located by matching its own label text against `mis_catalog`;
   a row whose label matches nothing is reported `ignored` /
   `unmatched_row`, never coerced onto the nearest-looking key.

2. **Validate, don't coerce; real numbers, real nulls, never a raw
   string.** A cell is one of three things: blank/placeholder (a real
   `None`, still "matched" — the document genuinely has nothing there),
   parseable (a real `float`/`int`), or garbled (`ambiguous`, the raw text
   preserved, no value produced). Placeholder detection matters as much as
   number parsing here: both source templates ship with every content cell
   holding literal instructional text ("[  ]", "[one-line commentary]",
   "e.g. ...") — treating that example text as if it were the founder's
   real answer would be the exact silent-overwrite failure this whole
   import/review split exists to prevent.

3. **Derive, never persist.** `vs Last Mo` (metrics), `Gap` (financials
   needs) and `Net Change` (headcount) are read from the document ONLY to
   cross-check against what this module itself would compute from the
   sibling cells actually present (Gap) or a caller-supplied previous
   period's own values (vs Last Mo, Net Change) — never stored, and a
   mismatch is reported as a `CrossCheckMismatch`, not silently accepted or
   silently dropped.

**Deliberately not attempted:** the quarterly template's seven "one
paragraph per entity" sections (IP Register, Collaborations, Publications,
Products, External Funding, and both halves of Milestone Review) fold 8-13
structured fields into a single free-running sentence with no per-field
textual anchor ("e.g. Patent — 'Novel Doppler algorithm...' (Healthcare,
National, IN2026XXXXXX, filed Feb-2026)."). Decomposing that into
`title`/`tech_area`/`filing_year`/... without an LLM is guessing, not
parsing — the exact failure mode "never invent" and "low-confidence fields
are returned blank and flagged" exist to forbid. Every non-blank bullet in
one of those sections is instead surfaced as a single `ambiguous` note
carrying the raw text, and produces zero rows in `entries` — a founder
reviewing the import sees "3 bullets found here, add them yourself" rather
than a table quietly populated with wrong guesses. The exception is
9.1 "Milestone Review" (`planned_vs_actual`) and monthly's Risks/Asks,
whose bullets DO carry literal per-field anchors in the source text
("Impact:", "Mitigation:", "Reason:", "Corrective action:", a leading
"N."), so those three are decomposed for real — see `_parse_risk_bullet`,
`_parse_ask_bullet`, `_parse_planned_vs_actual_bullet`.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import mis_catalog as cat

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class MisParseError(RuntimeError):
    """`code` is a short sentinel the router turns into a 4xx response."""

    def __init__(self, code: str, detail: str | None = None) -> None:
        super().__init__(detail or code)
        self.code = code
        self.detail = detail


@dataclass
class FieldNote:
    path: str
    reason: str
    raw: str | None = None


@dataclass
class CrossCheckMismatch:
    path: str
    doc_value: float | None
    computed_value: float | None


@dataclass
class ParsedMisDocument:
    kind: str
    narrative: dict[str, str] = field(default_factory=dict)
    metrics: list[dict] = field(default_factory=list)
    financials: list[dict] = field(default_factory=list)
    headcount: list[dict] = field(default_factory=list)
    entries: dict[str, list[dict]] = field(default_factory=dict)
    matched: list[str] = field(default_factory=list)
    ambiguous: list[FieldNote] = field(default_factory=list)
    ignored: list[FieldNote] = field(default_factory=list)
    cross_check_mismatches: list[CrossCheckMismatch] = field(default_factory=list)


# ── text normalisation + placeholder detection ───────────────────────────

_CURLY_QUOTES = {"‘": "'", "’": "'", "“": '"', "”": '"', "–": "-", "—": "-"}


def _normalise(text: str) -> str:
    t = (text or "").strip()
    for src, dst in _CURLY_QUOTES.items():
        t = t.replace(src, dst)
    t = re.sub(r"\s+", " ", t)
    return t.lower()


_BRACKET_RE = re.compile(r"^\[.*\]$")
_TRAILING_PAREN_RE = re.compile(r"\s*\([^)]*\)\s*$")


def _strip_trailing_paren(text: str) -> str:
    return _TRAILING_PAREN_RE.sub("", text).strip()


def _is_placeholder(raw: str | None) -> bool:
    """True for the template's own unfilled instructional content: an
    empty cell, a bracketed placeholder ("[  ]", "[one-line commentary]"),
    "e.g. ..." example prose, or a leftover "Add ..." instruction line
    ("Add one bullet per active collaboration."). Any of these appearing in
    an uploaded document means the founder has not actually answered that
    field yet -- reading it as their real answer is exactly the
    silent-wrong-data failure this parser exists to avoid."""
    t = (raw or "").strip()
    if not t:
        return True
    if _BRACKET_RE.match(t):
        return True
    low = t.lower()
    if low.startswith("e.g.") or low.startswith("e.g "):
        return True
    if low.startswith("add "):
        return True
    return False


def _clean_text(raw: str) -> str | None:
    """A free-text cell/value: `None` for a placeholder, else the trimmed
    text -- never the placeholder's own literal string."""
    if _is_placeholder(raw):
        return None
    return raw.strip() or None


def _parse_number(raw: str) -> tuple[float | None, bool]:
    """(value, ok). `ok=True` with `value=None` means genuinely blank
    (matched, nothing to import). `ok=False` means present-but-unparseable
    (ambiguous) -- the caller must not use `value` in that case."""
    if _is_placeholder(raw):
        return None, True
    cleaned = raw.strip().replace(",", "")
    cleaned = re.sub(r"^[₹$]\s*", "", cleaned)
    try:
        return float(cleaned), True
    except ValueError:
        return None, False


def _parse_int(raw: str) -> tuple[int | None, bool]:
    value, ok = _parse_number(raw)
    if not ok:
        return None, False
    if value is None:
        return None, True
    if not float(value).is_integer():
        return None, False
    return int(value), True


# ── kind detection + heading matching ────────────────────────────────────
#
# A catalog section title is not always the doc's FULL heading text: several
# quarterly headings carry a suffix the catalog's own (shorter) title omits
# -- "6. Financials — Deep Dive" (catalog title: "Financials"), "8. People —
# Full Snapshot" (catalog title: "People"), "2. IP Register (This Quarter +
# Cumulative)" (catalog title: "IP Register"). Matching is therefore
# startswith, not equality, picking the LONGEST catalog title that prefixes
# the heading -- longest-first so a short title can never shadow a more
# specific (longer) one that also matches.

def _catalog_titles_by_length(kind: str) -> list[tuple[str, str]]:
    return sorted(
        ((_normalise(s["title"]), s["id"]) for s in cat.SECTIONS[kind]),
        key=lambda pair: -len(pair[0]),
    )


def _match_heading(kind: str, heading_text: str) -> str | None:
    title = _normalise(re.sub(r"^\d+\.\s*", "", heading_text))
    for cat_title, section_id in _catalog_titles_by_length(kind):
        if title == cat_title or title.startswith(cat_title):
            return section_id
    return None


def _detect_kind(doc: Document) -> str | None:
    monthly_hits = quarterly_hits = 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else None
        if style != "Heading 1":
            continue
        if _match_heading("monthly", p.text):
            monthly_hits += 1
        if _match_heading("quarterly", p.text):
            quarterly_hits += 1
    if monthly_hits == 0 and quarterly_hits == 0:
        return None
    return "monthly" if monthly_hits >= quarterly_hits else "quarterly"


# ── narrative label maps (matched against the REAL document's own label
# text, not mis_catalog's prompt strings -- several prompts carry a
# section-number prefix or shorter wording than what the template actually
# prints; see the module docstring in mis_catalog for why prompts there are
# sometimes humanised rather than verbatim) ──────────────────────────────

_NARRATIVE_LABEL_MAP: dict[tuple[str, str], list[tuple[str, str]]] = {
    ("monthly", "exec_summary"): [
        ("headline win", "exec.headline_win"),
        ("biggest concern", "exec.biggest_concern"),
        ("commercial", "exec.commercial"),
        ("cash", "exec.cash"),
        ("top ask from artpark this month", "exec.top_ask"),
    ],
    ("monthly", "traction"): [
        ("active paid pilots", "traction.active_pilots"),
        ("conversions this month", "traction.conversions"),
        ("pipeline", "traction.pipeline"),
        ("losses", "traction.losses"),
        ("sharpest wedge", "traction.sharpest_wedge"),
        ("what isn't working", "traction.not_working"),
    ],
    ("monthly", "team_hiring"): [
        ("headcount", "team.headcount"),
        ("joined this month", "team.joined"),
        ("left this month", "team.left"),
        ("open roles", "team.open_roles"),
    ],
    ("monthly", "financials_fundraising"): [
        ("revenue", "fin.revenue"),
        ("gross margin", "fin.gross_margin"),
        ("cash burn", "fin.cash_burn"),
        ("cash balance & runway", "fin.cash_and_runway"),
        ("round in progress", "fund.round_in_progress"),
        ("investor conversations", "fund.investor_conversations"),
        ("non-dilutive capital", "fund.non_dilutive"),
    ],
    ("monthly", "happy_news"): [
        ("field story", "happy.field_story"),
        ("recognition", "happy.recognition"),
        ("demos & links", "happy.demos_links"),
    ],
    ("quarterly", "glance"): [
        ("strategic theme of the quarter", "glance.strategic_theme"),
        ("biggest milestone achieved", "glance.biggest_milestone"),
        ("biggest miss or setback", "glance.biggest_miss"),
        ("commercial + funding position", "glance.commercial_funding_position"),
        ("what we're betting on next quarter", "glance.next_quarter_bet"),
    ],
    ("quarterly", "financials"): [
        ("cash in bank", "fin6.cash_in_bank"),
        ("quarterly burn", "fin6.quarterly_burn"),
        ("runway", "fin6.runway"),
        ("gross margin trajectory", "fin6.gross_margin_trajectory"),
        ("gap-closing plan", "fin6.gap_closing_plan"),
        ("revenue commentary", "fin6.revenue_commentary"),
    ],
    ("quarterly", "headcount"): [
        ("diversity snapshot", "people.diversity"),
        ("key hires this quarter", "people.key_hires"),
        ("attrition", "people.attrition"),
        ("team structure changes", "people.structure_changes"),
    ],
}

# section id -> narrative field id, for a subheading zone that carries the
# answer as plain bullets with NO "Label:" prefix at all (quarterly 9.3).
_SINGLE_FIELD_NARRATIVE_TOKENS: dict[str, tuple[str, str]] = {
    "9.3": ("planned_vs_actual", "gc.strategic_questions"),
}

# quarterly subheading token ("2.1", "3.2", ...) -> bucket, per
# mis-templates.md's own numbering. Monthly's 4.1/4.2/8.1/8.2 subheadings
# are deliberately NOT here: monthly narrative sections are scanned as one
# whole Heading-1 zone regardless of subheading (see _handle_bullet).
_BUCKET_TOKEN_MAP: dict[str, str] = {
    "2.1": "filed", "2.2": "granted", "2.3": "rejected",
    "2.4": "international", "2.5": "cumulative",
    "3.1": "active", "3.2": "new", "3.3": "completed", "3.4": "in_discussion",
    "4.1": "published", "4.2": "in_review", "4.3": "standards_policy",
}
# section ids that are pure free-prose "one bullet per entity" -- never
# decomposed into fields, see module docstring.
_PROSE_ONLY_SECTIONS = {"ip_assets", "collaborations", "publications", "products", "funding"}

_SUBHEADING_TOKEN_RE = re.compile(r"^(\d+\.\d+)")


def _subheading_token(text: str) -> str | None:
    m = _SUBHEADING_TOKEN_RE.match(text.strip())
    return m.group(1) if m else None


def _match_label(label_part: str, candidates: list[tuple[str, str]]) -> str | None:
    norm = _normalise(label_part)
    for label, field_id in candidates:
        if norm.startswith(label):
            return field_id
    return None


# ── risks / asks / planned_vs_actual: bullets with real embedded anchors ──

_RISK_HEAD_RE = re.compile(r"^Risk\s*\d+\s*\(([^)]*)\)\s*:\s*(.*)$", re.IGNORECASE | re.DOTALL)
_IMPACT_RE = re.compile(r"\bImpact:\s*", re.IGNORECASE)
_MITIGATION_RE = re.compile(r"\bMitigation:\s*", re.IGNORECASE)


def _parse_risk_bullet(text: str, idx: int) -> tuple[dict | None, list[FieldNote]]:
    m = _RISK_HEAD_RE.match(text.strip())
    if not m:
        return None, []
    sev_raw, body = m.groups()
    if _is_placeholder(body):
        return None, []
    notes: list[FieldNote] = []
    row: dict = {}
    sev_low = sev_raw.lower()
    if "red" in sev_low:
        row["severity"] = "red"
    elif "amber" in sev_low:
        row["severity"] = "amber"
    else:
        notes.append(FieldNote(path=f"entries.risks[{idx}].severity", reason="invalid_value", raw=sev_raw))

    mit_parts = _MITIGATION_RE.split(body, maxsplit=1)
    head, mitigation = mit_parts[0], (mit_parts[1].strip() if len(mit_parts) > 1 else None)
    imp_parts = _IMPACT_RE.split(head, maxsplit=1)
    what_happened, impact = imp_parts[0].strip(), (imp_parts[1].strip() if len(imp_parts) > 1 else None)

    if what_happened:
        row["what_happened"] = what_happened
    if impact:
        row["impact"] = impact
    if mitigation:
        row["mitigation"] = mitigation
    return row, notes


_ASK_RE = re.compile(r"^(\d+)\.\s*(.*)$", re.DOTALL)


def _parse_ask_bullet(text: str, idx: int) -> tuple[dict | None, list[FieldNote]]:
    m = _ASK_RE.match(text.strip())
    if not m:
        return None, []
    num_raw, rest = m.groups()
    if _is_placeholder(rest):
        return None, []
    row = {"priority": int(num_raw), "ask": rest.strip()}
    notes = [FieldNote(path=f"entries.asks[{idx}].category", reason="no_textual_anchor", raw=None)]
    return row, notes


_PLANNED_RE = re.compile(r"^\s*Planned:\s*", re.IGNORECASE)
_ACHIEVED_RE = re.compile(r"\b(?:Achieved|Actual):\s*", re.IGNORECASE)
_REASON_RE = re.compile(r"\bReason:\s*", re.IGNORECASE)
_CORRECTIVE_RE = re.compile(r"\bCorrective action:\s*", re.IGNORECASE)


def _parse_planned_vs_actual_bullet(text: str, idx: int) -> tuple[dict | None, list[FieldNote]]:
    t = text.strip()
    if _is_placeholder(t):
        return None, []
    if not _PLANNED_RE.match(t):
        return None, [FieldNote(path=f"entries.planned_vs_actual[{idx}]",
                                 reason="freeform_entry_not_decomposed", raw=t)]
    rest = _PLANNED_RE.sub("", t, count=1)
    corrective_parts = _CORRECTIVE_RE.split(rest, maxsplit=1)
    head1 = corrective_parts[0]
    corrective = corrective_parts[1].strip() if len(corrective_parts) > 1 else None
    reason_parts = _REASON_RE.split(head1, maxsplit=1)
    head2 = reason_parts[0]
    reason = reason_parts[1].strip() if len(reason_parts) > 1 else None
    achieved_parts = _ACHIEVED_RE.split(head2, maxsplit=1)
    planned = achieved_parts[0].strip()
    achieved = achieved_parts[1].strip() if len(achieved_parts) > 1 else None

    row: dict = {}
    if planned:
        row["planned"] = planned
    if achieved:
        row["achieved"] = achieved
    if reason:
        row["reason"] = reason
    if corrective:
        row["corrective_action"] = corrective
    return row, []


def _parse_next_milestone_bullet(text: str) -> dict | None:
    if _is_placeholder(text):
        return None
    return {"milestone": text.strip()}


# ── table parsing: monthly Key Metrics ────────────────────────────────────

def _parse_metrics_table(
    table: Table, prev_actual_by_metric_key: dict[str, float | None] | None,
) -> tuple[list[dict], list[FieldNote], list[FieldNote], list[CrossCheckMismatch]]:
    rows_out: list[dict] = []
    ambiguous: list[FieldNote] = []
    ignored: list[FieldNote] = []
    cross_checks: list[CrossCheckMismatch] = []
    group_labels = {g["label"] for g in cat.METRIC_GROUPS}
    metric_by_label = {_normalise(m["label"]): m for m in cat.METRICS}
    prev_actual_by_metric_key = prev_actual_by_metric_key or {}

    for ridx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if ridx == 0 or len(cells) < 5:
            continue
        if len(set(cells)) == 1 and cells[0] in group_labels:
            continue  # group-heading row, structural, not data

        m = metric_by_label.get(_normalise(cells[0]))
        if m is None:
            ignored.append(FieldNote(path=f"metrics[{ridx}]", reason="unmatched_row", raw=cells[0]))
            continue
        key = m["key"]

        target, t_ok = _parse_number(cells[1])
        if not t_ok:
            ambiguous.append(FieldNote(path=f"metrics.{key}.target", reason="unparseable_number", raw=cells[1]))
            target = None

        actual, a_ok = _parse_number(cells[2])
        if not a_ok:
            ambiguous.append(FieldNote(path=f"metrics.{key}.actual", reason="unparseable_number", raw=cells[2]))
            actual = None

        # vs Last Mo (col 3) is read only to cross-check against actual -
        # prev_actual, supplied by the caller (mis_query already computes
        # this for the period being imported into) -- never persisted.
        vs_last_doc, vl_ok = _parse_number(cells[3])
        if vl_ok and vs_last_doc is not None and actual is not None:
            prev = prev_actual_by_metric_key.get(key)
            if prev is not None:
                computed = actual - prev
                if abs(computed - vs_last_doc) > 1e-9:
                    cross_checks.append(CrossCheckMismatch(
                        path=f"metrics.vs_last.{key}", doc_value=vs_last_doc, computed_value=computed,
                    ))

        commentary = _clean_text(cells[4])

        if key == "trl_level" and actual is not None:
            # Server-computed from verified AIR -- a founder-supplied value
            # here is read only to note that it was discarded, never kept.
            ignored.append(FieldNote(path="metrics.trl_level.actual", reason="computed_field", raw=cells[2]))
            actual = None

        rows_out.append({"metric_key": key, "target": target, "actual": actual, "commentary": commentary})

    return rows_out, ambiguous, ignored, cross_checks


# ── table parsing: monthly Milestones ─────────────────────────────────────

def _parse_milestones_table(table: Table) -> tuple[list[dict], list[FieldNote]]:
    rows_out: list[dict] = []
    ambiguous: list[FieldNote] = []
    fields = {f["key"]: f for f in cat.entry_fields("milestones")}
    order = ["milestone", "owner", "status", "notes"]

    for ridx, row in enumerate(table.rows):
        if ridx == 0:
            continue
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 4:
            continue
        row_out: dict = {}
        for key, text in zip(order, cells):
            f = fields[key]
            if _is_placeholder(text):
                continue
            if f["type"] == "choice":
                opt = next((o for o in f["options"] if o.lower() == text.strip().lower()), None)
                if opt is None:
                    ambiguous.append(FieldNote(
                        path=f"entries.milestones[{len(rows_out)}].{key}",
                        reason="invalid_value", raw=text,
                    ))
                else:
                    row_out[key] = opt
            else:
                row_out[key] = text.strip()
        if row_out:
            rows_out.append(row_out)

    return rows_out, ambiguous


# ── table parsing: quarterly Financials (6.1 annual revenue, 6.2 needs) ──

_ANNUAL_REVENUE_ROW_LABELS = {
    _normalise(_strip_trailing_paren(s["label"])): s["key"]
    for s in cat.FINANCIAL_SERIES["annual_revenue"]
}
_NEEDS_ROW_LABELS = {_normalise(s["label"]): s["key"] for s in cat.FINANCIAL_SERIES["needs"]}


def _classify_financial_table(table: Table) -> str | None:
    header = " ".join(c.text.strip() for c in table.rows[0].cells).upper()
    if "FY" in header:
        return "annual_revenue"
    if "Q1" in header:
        return "needs"
    return None


def _parse_annual_revenue_table(table: Table, fy_start_year: int) -> tuple[list[dict], list[FieldNote]]:
    buckets = cat.annual_revenue_buckets(fy_start_year)
    rows_out: list[dict] = []
    ignored: list[FieldNote] = []
    for ridx, row in enumerate(table.rows):
        if ridx == 0:
            continue
        cells = [c.text.strip() for c in row.cells]
        series = _ANNUAL_REVENUE_ROW_LABELS.get(_normalise(_strip_trailing_paren(cells[0])))
        if series is None:
            ignored.append(FieldNote(path=f"financials[{ridx}]", reason="unmatched_row", raw=cells[0]))
            continue
        # Bucket columns are matched POSITIONALLY, not by the doc's own
        # header text: the template hard-codes literal FY labels
        # ("FY 21-22" .. "FY 25-26 Proj.") that go stale the moment a
        # fiscal year turns over -- mis_catalog.annual_revenue_buckets'
        # own docstring is explicit these must be treated as relative to
        # the importing period's fiscal year, not read as literal text.
        for col_idx in range(1, min(len(cells), len(buckets) + 1)):
            value, ok = _parse_number(cells[col_idx])
            if not ok:
                value = None  # unparseable annual-revenue cells are rare/edge; treat as blank rather than block the row
            rows_out.append({"series": series, "bucket": buckets[col_idx - 1], "amount": value})
    return rows_out, ignored


def _parse_needs_table(table: Table) -> tuple[list[dict], list[FieldNote], list[CrossCheckMismatch]]:
    header_cells = [c.text.strip() for c in table.rows[0].cells]
    buckets = cat.FINANCIAL_BUCKETS["needs"]
    col_bucket: dict[int, str] = {}
    for col_idx in range(1, len(header_cells)):
        text = re.sub(r"\s+", " ", header_cells[col_idx].strip())
        col_bucket[col_idx] = text if text in buckets else (
            buckets[col_idx - 1] if col_idx - 1 < len(buckets) else text
        )

    rows_out: list[dict] = []
    ignored: list[FieldNote] = []
    doc_gap_by_bucket: dict[str, float | None] = {}
    values_by_bucket: dict[str, dict[str, float | None]] = {b: {} for b in buckets}

    for ridx, row in enumerate(table.rows):
        if ridx == 0:
            continue
        cells = [c.text.strip() for c in row.cells]
        series = _NEEDS_ROW_LABELS.get(_normalise(cells[0]))
        if series is None:
            ignored.append(FieldNote(path=f"financials[{ridx}]", reason="unmatched_row", raw=cells[0]))
            continue
        for col_idx, bucket in col_bucket.items():
            if col_idx >= len(cells):
                continue
            value, ok = _parse_number(cells[col_idx])
            if not ok:
                value = None
            if series == "needs_gap":
                doc_gap_by_bucket[bucket] = value
                continue  # computed-only: never persisted
            rows_out.append({"series": series, "bucket": bucket, "amount": value})
            values_by_bucket.setdefault(bucket, {})[series] = value

    cross_checks: list[CrossCheckMismatch] = []
    for bucket, doc_gap in doc_gap_by_bucket.items():
        if doc_gap is None:
            continue
        vals = values_by_bucket.get(bucket, {})
        total, confirmed, projected = vals.get("needs_total"), vals.get("needs_confirmed"), vals.get("needs_projected")
        if total is None or confirmed is None or projected is None:
            continue
        computed = total - confirmed - projected
        if abs(computed - doc_gap) > 1e-9:
            cross_checks.append(CrossCheckMismatch(
                path=f"financials.needs_gap.{bucket}", doc_value=doc_gap, computed_value=computed,
            ))

    return rows_out, ignored, cross_checks


# ── table parsing: quarterly Headcount ────────────────────────────────────

_HEADCOUNT_ROW_LABELS = {_normalise(c["label"]): c["key"] for c in cat.HEADCOUNT_CATEGORIES}


def _parse_headcount_table(
    table: Table, prev_headcount_current: dict[str, float | None] | None,
) -> tuple[list[dict], list[FieldNote], list[CrossCheckMismatch]]:
    rows_out: list[dict] = []
    ignored: list[FieldNote] = []
    cross_checks: list[CrossCheckMismatch] = []
    prev_headcount_current = prev_headcount_current or {}

    for ridx, row in enumerate(table.rows):
        if ridx == 0:
            continue
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 5:
            continue
        if _normalise(cells[0]) == "total":
            continue  # structural aggregate row, never a category
        category = _HEADCOUNT_ROW_LABELS.get(_normalise(cells[0]))
        if category is None:
            ignored.append(FieldNote(path=f"headcount[{ridx}]", reason="unmatched_row", raw=cells[0]))
            continue

        current_count, cc_ok = _parse_int(cells[1])
        exited, ex_ok = _parse_int(cells[2])
        # Net Change (col 3) is read only to cross-check against
        # current_count - the PREVIOUS quarterly period's own current_count
        # (a stock-over-time delta) -- never current_count - exited, the
        # sign-error bug already fixed in mis_query.py. Never persisted.
        net_change_doc, nc_ok = _parse_number(cells[3])
        if nc_ok and net_change_doc is not None and cc_ok and current_count is not None:
            prev = prev_headcount_current.get(category)
            if prev is not None:
                computed = current_count - prev
                if abs(computed - net_change_doc) > 1e-9:
                    cross_checks.append(CrossCheckMismatch(
                        path=f"headcount.net_change.{category}",
                        doc_value=net_change_doc, computed_value=computed,
                    ))
        remarks = _clean_text(cells[4])

        rows_out.append({
            "category": category,
            "current_count": current_count if cc_ok else None,
            "exited": exited if ex_ok else None,
            "remarks": remarks,
        })

    return rows_out, ignored, cross_checks


# ── the document walk ─────────────────────────────────────────────────────

def _walk(
    doc: Document, kind: str, fy_start_year: int | None,
    prev_actual_by_metric_key: dict[str, float | None] | None,
    prev_headcount_current: dict[str, float | None] | None,
) -> ParsedMisDocument:
    result = ParsedMisDocument(kind=kind)
    current_section: str | None = None
    current_token: str | None = None
    gc_bullets: list[str] = []
    entry_counters: dict[str, int] = {}

    def _next_idx(section_id: str) -> int:
        idx = entry_counters.get(section_id, 0)
        entry_counters[section_id] = idx + 1
        return idx

    def _handle_bullet(text: str) -> None:
        nonlocal current_token
        if current_section is None:
            return

        # quarterly 9.3 -- freeform bullets, no label, single narrative field
        if kind == "quarterly" and current_section == "planned_vs_actual" and current_token == "9.3":
            if not _is_placeholder(text):
                gc_bullets.append(text.strip())
            return

        label_map = _NARRATIVE_LABEL_MAP.get((kind, current_section))
        if label_map:
            label_part, sep, rest = text.partition(":")
            if sep:
                field_id = _match_label(label_part, label_map)
                if field_id:
                    value = _clean_text(rest)
                    if value is not None:
                        result.narrative[field_id] = value
                        result.matched.append(f"narrative.{field_id}")
                    return

        section_type = cat.section(kind, current_section)["type"] if _has_section(kind, current_section) else None

        if kind == "monthly" and current_section == "risks":
            row, notes = _parse_risk_bullet(text, _peek_idx("risks"))
            result.ambiguous.extend(notes)
            if row:
                result.entries.setdefault("risks", []).append(row)
                _next_idx("risks")
            return

        if kind == "monthly" and current_section == "asks":
            row, notes = _parse_ask_bullet(text, _peek_idx("asks"))
            result.ambiguous.extend(notes)
            if row:
                result.entries.setdefault("asks", []).append(row)
                _next_idx("asks")
            return

        if kind == "quarterly" and current_section == "planned_vs_actual" and current_token != "9.2":
            row, notes = _parse_planned_vs_actual_bullet(text, _peek_idx("planned_vs_actual"))
            result.ambiguous.extend(notes)
            if row:
                result.entries.setdefault("planned_vs_actual", []).append(row)
                _next_idx("planned_vs_actual")
            return

        if kind == "quarterly" and current_section == "planned_vs_actual" and current_token == "9.2":
            row = _parse_next_milestone_bullet(text)
            if row:
                result.entries.setdefault("next_milestones", []).append(row)
            return

        if kind == "quarterly" and current_section in _PROSE_ONLY_SECTIONS:
            idx = _next_idx(current_section)
            if not _is_placeholder(text):
                result.ambiguous.append(FieldNote(
                    path=f"entries.{current_section}[{idx}]",
                    reason="freeform_entry_not_decomposed", raw=text.strip(),
                ))
            return

        # Anything else non-blank under a recognised section that matched no
        # handler above is genuinely unrecognised content -- surfaced, not
        # silently dropped.
        if not _is_placeholder(text):
            result.ignored.append(FieldNote(
                path=f"{current_section}.unrecognized_bullet", reason="unrecognized_content", raw=text.strip(),
            ))

    def _peek_idx(section_id: str) -> int:
        return entry_counters.get(section_id, 0)

    def _handle_table(table: Table) -> None:
        if kind == "monthly" and current_section == "key_metrics":
            rows, amb, ign, cc = _parse_metrics_table(table, prev_actual_by_metric_key)
            result.metrics.extend(rows)
            result.ambiguous.extend(amb)
            result.ignored.extend(ign)
            result.cross_check_mismatches.extend(cc)
            result.matched.extend(f"metrics.{r['metric_key']}" for r in rows)
            return
        if kind == "monthly" and current_section == "milestones":
            rows, amb = _parse_milestones_table(table)
            result.entries.setdefault("milestones", []).extend(rows)
            result.ambiguous.extend(amb)
            return
        if kind == "quarterly" and current_section == "financials":
            table_kind = _classify_financial_table(table)
            if table_kind == "annual_revenue":
                rows, ign = _parse_annual_revenue_table(table, fy_start_year)  # type: ignore[arg-type]
                result.financials.extend(rows)
                result.ignored.extend(ign)
            elif table_kind == "needs":
                rows, ign, cc = _parse_needs_table(table)
                result.financials.extend(rows)
                result.ignored.extend(ign)
                result.cross_check_mismatches.extend(cc)
            return
        if kind == "quarterly" and current_section == "headcount":
            rows, ign, cc = _parse_headcount_table(table, prev_headcount_current)
            result.headcount.extend(rows)
            result.ignored.extend(ign)
            result.cross_check_mismatches.extend(cc)
            return

    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            style = item.style.name if item.style else None
            text = item.text.strip()
            if not text:
                continue
            if style == "Heading 1":
                current_section = _match_heading(kind, text)
                current_token = None
                continue
            if style == "Heading 2":
                current_token = _subheading_token(text)
                continue
            if style != "List Paragraph":
                continue
            _handle_bullet(text)
        elif isinstance(item, Table):
            _handle_table(item)

    if gc_bullets:
        result.narrative["gc.strategic_questions"] = "\n".join(gc_bullets)
        result.matched.append("narrative.gc.strategic_questions")

    return result


def _has_section(kind: str, section_id: str) -> bool:
    return any(s["id"] == section_id for s in cat.SECTIONS[kind])


# ── top-level entry point ─────────────────────────────────────────────────

def parse_mis_document(
    *,
    file_bytes: bytes,
    mime: str,
    fy_start_year: int | None = None,
    prev_actual_by_metric_key: dict[str, float | None] | None = None,
    prev_headcount_current: dict[str, float | None] | None = None,
) -> ParsedMisDocument:
    """Parse an uploaded MIS `.docx` into `mis_catalog`'s own shape.

    `fy_start_year` is required for a quarterly document (the calendar year
    the importing period's own fiscal year starts in) -- it is what makes
    the annual-revenue grid's positional bucket mapping correct; there is no
    safe default (no fail-open defaults), so a caller that omits it for a
    quarterly upload gets a loud `fy_start_year_required` rather than a
    silently wrong bucket assignment.

    `prev_actual_by_metric_key`/`prev_headcount_current` are optional caller
    context (typically the SAME previous-period values `mis_query.
    period_bundle`'s own `derived` block already computed) used only to
    cross-check the document's own `vs Last Mo`/`Net Change` columns -- see
    the module docstring's ruling 3. Omitting them just means those two
    cross-checks are skipped, not that anything is guessed.
    """
    if not file_bytes:
        raise MisParseError("empty_document", "Empty file uploaded.")
    mime = (mime or "").lower().strip()
    if mime != DOCX_MIME:
        raise MisParseError(
            "unsupported_mime",
            "Please upload the filled MIS template as .docx.",
        )

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise MisParseError("empty_document", f"Could not read .docx: {exc}") from exc

    kind = _detect_kind(doc)
    if kind is None:
        raise MisParseError(
            "unrecognised_template",
            "This doesn't look like an ARTPARK monthly update or quarterly review template.",
        )
    if kind == "quarterly" and fy_start_year is None:
        raise MisParseError(
            "fy_start_year_required",
            "The importing period's fiscal year is required to parse a quarterly review.",
        )

    return _walk(doc, kind, fy_start_year, prev_actual_by_metric_key, prev_headcount_current)
