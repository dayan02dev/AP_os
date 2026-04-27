"""Offline-template parsing pipeline.

Stages:
  1. Extract per-question raw text via anchor markers
       >>> ANSWER Q9 START >>>
       (applicant text)
       <<< ANSWER Q9 END <<<
     For .docx we walk the python-docx Document tree (paragraphs + tables);
     for .pdf we use the same pypdf path the resume parser uses.
  2. Read Word checkbox state for the two MCQ questions (Q10, Q14) directly
     from the raw word/document.xml — python-docx doesn't expose w14:checkbox
     content controls so we crack the .docx zip ourselves and walk the SDT
     elements in document order, matching their positional index against
     the option list inside each Q's anchor block.
  3. Hand the extracted dict {Q9..Q19} to llm_service.normalize_template_answers
     for whitespace tidy + MCQ normalisation + empty-detection (Gemini Flash).
  4. Caller persists `parsed_data` on the application_templates row.

Failure modes (all stamp parse_status='failed', parse_error=<sentinel>):
  - empty_document          — file has zero usable text
  - no_anchors_detected     — applicant deleted the markers
  - unsupported_mime        — caller already screened, but defence in depth
  - llm_normalization_failed — Gemini round-trip failed after retries
"""

from __future__ import annotations

import io
import logging
import re
import zipfile
from typing import Any

from docx import Document
from pypdf import PdfReader

from .llm_service import LLMParseError, OpenRouterClient

log = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

# Anchor markers — the literal strings the template ships with. Allow
# tolerant whitespace inside the marker but require the Q-number tag.
_ANCHOR_RE = re.compile(
    r">>>\s*ANSWER\s+(Q\d+)\s+START\s*>>>(.*?)<<<\s*ANSWER\s+\1\s+END\s*<<<",
    re.DOTALL,
)

# Per-cell intro line + option lines. The template ships with:
#     "Click the tick box next to your choice. Select only one."
#     "    A.  Yes"
#     "    B.  No"
# We strip the intro and read each "<letter>.  <label>" line.
_OPTION_LINE_RE = re.compile(r"^\s*([A-Z])\.\s+(.+?)\s*$")
_TICK_INSTRUCTIONS_RE = re.compile(
    r"click the tick box next to your choice\.?\s*select only one\.?",
    re.IGNORECASE,
)

# Word's content-control checkbox lives in the w14 namespace; we don't
# bother with proper XML parsing for a one-attribute lookup.
_CHECKBOX_RE = re.compile(
    r"<w14:checkbox>.*?<w14:checked\s+w14:val=\"([01])\"\s*/?>",
    re.DOTALL,
)

# Question IDs we expect; ordered for stable iteration.
QUESTION_IDS = [f"Q{i}" for i in range(9, 20)]
# Q10 + Q14 are the two checkbox/MCQ questions — every other Q is a long-text
# essay. The ordering of MCQs in document order matters: Q10 ships its
# checkboxes first, then Q14, so positional checkbox-XML reads align.
MCQ_QUESTIONS = ("Q10", "Q14")


class TemplateParseError(RuntimeError):
    """Raised when parsing or LLM normalisation cannot complete.

    `code` is a short sentinel (e.g. "empty_document") suitable for
    persistence in `application_templates.parse_error` and for the
    frontend to switch on.
    """

    def __init__(self, code: str, detail: str | None = None) -> None:
        super().__init__(detail or code)
        self.code = code
        self.detail = detail


# ── DOCX-specific extraction ──────────────────────────────────────────────

def _docx_concatenated_text(file_bytes: bytes) -> str:
    """Walk the python-docx tree once and emit a single text blob.

    We need both paragraphs (the question prose) and table cells (the
    answer bodies are inside single-cell tables) in their natural order.
    `Document.element.body.iter()` would mix content with proofing/run
    XML — easier to use the document's `.paragraphs` + `.tables` since the
    template never nests tables inside cells.
    """
    doc = Document(io.BytesIO(file_bytes))
    chunks: list[str] = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(c for c in chunks if c is not None)


def _docx_checkbox_states(file_bytes: bytes) -> list[bool]:
    """Return checkbox states in document order.

    Each Word checkbox content control encodes `w14:val="0"` or `"1"`. We
    read the ZIP's `word/document.xml` directly and regex out the value
    list. With Q10 (2 boxes) + Q14 (7 boxes) = 9 entries on a fresh
    template; the order is stable as long as the applicant doesn't
    reorder questions (which the template explicitly tells them not to).
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            with zf.open("word/document.xml") as fh:
                xml = fh.read().decode("utf-8", errors="replace")
    except (zipfile.BadZipFile, KeyError) as exc:
        # Treat as no checkboxes; the LLM will still try to read text-only
        # Q10/Q14 answers (e.g. applicant typed "Yes" in the cell).
        log.warning("could not read document.xml for checkbox state: %s", exc)
        return []
    return [m.group(1) == "1" for m in _CHECKBOX_RE.finditer(xml)]


# ── PDF-specific extraction ───────────────────────────────────────────────

def _pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n".join(pages)


# ── Anchor-driven slicing ─────────────────────────────────────────────────

def _split_options_from_block(block: str) -> tuple[str, list[tuple[str, str]]]:
    """Pull option lines out of an MCQ cell.

    Returns (cleaned_freetext, [(letter, label), ...]).
    `cleaned_freetext` is whatever the applicant typed besides the option
    lines themselves (rare but possible — they may write 'A' on a fresh
    line). The LLM uses both signals.
    """
    options: list[tuple[str, str]] = []
    leftover: list[str] = []
    for raw_line in block.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if _TICK_INSTRUCTIONS_RE.search(line):
            continue
        m = _OPTION_LINE_RE.match(line)
        if m:
            options.append((m.group(1), m.group(2).strip()))
            continue
        leftover.append(line)
    return ("\n".join(leftover).strip(), options)


def _extract_anchor_blocks(text: str) -> dict[str, str]:
    """Run the anchor regex and return raw cell contents per question."""
    blocks: dict[str, str] = {}
    for m in _ANCHOR_RE.finditer(text):
        qid = m.group(1).upper()
        blocks[qid] = m.group(2).strip()
    return blocks


# ── Top-level entry point ─────────────────────────────────────────────────

async def parse_template(
    *,
    file_bytes: bytes,
    mime: str,
    user_id: str | None = None,
) -> dict[str, Any]:
    """Run the whole pipeline and return the normalised dict.

    Output shape (always emitted, never partial):
        {
          "Q9":  "...essay text..." | None,
          "Q10": "Yes" | "No" | None,
          ...
          "Q14": "Lab demos / proof of concept" | None,
          ...
          "Q19": "..." | None,
        }
    """
    if not file_bytes:
        raise TemplateParseError("empty_document", "Empty file uploaded.")

    mime = (mime or "").lower().strip()

    # 1. Extract concatenated text + checkbox states.
    if mime == DOCX_MIME:
        try:
            full_text = _docx_concatenated_text(file_bytes)
            checkbox_states = _docx_checkbox_states(file_bytes)
        except Exception as exc:
            log.warning("docx extraction failed", extra={"err": str(exc)})
            raise TemplateParseError("empty_document", f"Could not read .docx: {exc}") from exc
    elif mime == PDF_MIME:
        try:
            full_text = _pdf_text(file_bytes)
        except Exception as exc:
            log.warning("pdf extraction failed", extra={"err": str(exc)})
            raise TemplateParseError("empty_document", f"Could not read PDF: {exc}") from exc
        checkbox_states = []  # PDFs render checkboxes as glyphs; LLM handles
    else:
        raise TemplateParseError(
            "unsupported_mime",
            "Please upload the filled template as .docx (preferred) or .pdf.",
        )

    if not full_text.strip():
        raise TemplateParseError("empty_document", "No text could be extracted.")

    # 2. Split into per-question anchor blocks.
    blocks = _extract_anchor_blocks(full_text)

    # Fallback: if the applicant uploaded a doc without our anchor markers
    # (a re-typed copy, a Google-Docs export that mangled them, or a
    # different sample file), hand the whole document to Gemini and let it
    # locate each answer. The output schema is identical so the rest of
    # the pipeline doesn't care which path produced the dict. We require
    # at least 3 anchor matches before trusting the deterministic path —
    # 1–2 stray matches usually mean the markers are partially corrupted
    # and the freeform pass will give a better answer.
    if len(blocks) < 3:
        log.info(
            "template anchor extraction sparse, falling back to freeform LLM",
            extra={"user_id": user_id, "anchor_count": len(blocks)},
        )
        try:
            normalised = await OpenRouterClient().extract_template_answers_freeform(
                full_text, user_id=user_id,
            )
        except LLMParseError as exc:
            # If the LLM fallback also fails, surface the friendlier
            # "no_anchors_detected" message so the UI can advise the
            # applicant to re-download the template.
            if not blocks:
                raise TemplateParseError(
                    "no_anchors_detected",
                    "We couldn't find any of the answer markers in this file, "
                    "and the fallback parser couldn't extract answers either. "
                    "Please download the template above and fill answers between "
                    "the >>> ANSWER QN START >>> markers.",
                ) from exc
            raise TemplateParseError("llm_normalization_failed", str(exc)) from exc
        return {qid: normalised.get(qid) for qid in QUESTION_IDS}

    # 3. For Q10 / Q14, fold checkbox state in by position.
    #    Q10 has 2 options (Yes/No) → checkbox indices [0, 1]
    #    Q14 has 7 options          → checkbox indices [2..8]
    mcq_payload: dict[str, dict[str, Any]] = {}
    cb_cursor = 0
    for qid in MCQ_QUESTIONS:
        block = blocks.get(qid, "")
        leftover, options = _split_options_from_block(block)
        states_for_q: list[tuple[str, str, bool | None]] = []
        for letter, label in options:
            state: bool | None
            if cb_cursor < len(checkbox_states):
                state = checkbox_states[cb_cursor]
            else:
                state = None  # PDFs / older Word — let the LLM decide
            states_for_q.append((letter, label, state))
            cb_cursor += 1
        mcq_payload[qid] = {
            "free_text": leftover,
            "options": [
                {"letter": l, "label": lbl, "checked": st}
                for (l, lbl, st) in states_for_q
            ],
        }

    # 4. Build the LLM prompt input. Essay questions go through as raw text;
    #    MCQ questions come with their option grids so Gemini can return the
    #    exact label (or null if nothing is selected).
    payload: dict[str, Any] = {}
    for qid in QUESTION_IDS:
        if qid in MCQ_QUESTIONS:
            payload[qid] = mcq_payload.get(qid, {"free_text": "", "options": []})
        else:
            payload[qid] = {"free_text": blocks.get(qid, "").strip()}

    try:
        normalised = await OpenRouterClient().normalize_template_answers(
            payload, user_id=user_id,
        )
    except LLMParseError as exc:
        raise TemplateParseError("llm_normalization_failed", str(exc)) from exc

    # Always return the full key-set so callers can assume shape.
    return {qid: normalised.get(qid) for qid in QUESTION_IDS}


# Mapping consumed by the apply-to-application flow.
# Q → DB column on public.applications.
QUESTION_TO_APPLICATION_COLUMN: dict[str, str] = {
    "Q9":  "problem_describe",
    "Q10": "problem_defined",
    "Q11": "solution_describe",
    "Q12": "solution_core_tech",
    "Q13": "solution_contrarian_insight",
    "Q14": "solution_stage",
    "Q15": "execution_will_break",
    "Q16": "execution_milestone",
    "Q17": "execution_infrastructure",
    "Q18": "execution_failure",
    "Q19": "execution_hwsw_integration",
}
