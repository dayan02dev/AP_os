"""Mechanical .docx -> JSON template extractor.

Run BY A HUMAN when a new agreement .docx arrives; the output is committed
to the repo and is what the runtime loads (app/services/agreements.py never
opens a .docx — the production Lambda has no Word file on disk).

This script has NO knowledge of what any field means: it does not know which
[.] or [Named Token] is a founder's name versus ARTPARK's insurance limit.
That interpretation belongs entirely to app/services/agreements.py, which
reads the committed JSON for a specific agreement slug. Keeping this generic
is what makes it reusable for any future agreement.

Body children are walked via document.element.body.iterchildren() rather
than the higher-level document.paragraphs / document.tables properties.
Those flattened lists lose each table's position relative to the
surrounding paragraphs -- walking the raw body XML in order is the only way
to preserve real document order when paragraphs and tables interleave.

TRACKED CHANGES. A redlined .docx (e.g. a reviewer's accepted-but-not-yet-
"Accept All"-clicked edits) stores insertions inside <w:ins> and deletions
inside <w:del>/<w:delText>, both nested one level below <w:p> or <w:tc> --
one level deeper than the direct-child <w:r> runs python-docx's own
Paragraph.text / _Cell.text walk. That shallow walk silently DROPS both
sides of any tracked change: neither the old deleted text nor the newly
inserted replacement ever reaches .text, leaving orphaned fragments (e.g.
"having PAN s/o/d/o resident of" where "[•]"/"[PAN Number]" used to be).
_accepted_paragraph_text()/_accepted_cell_text() below fix this by collecting
every <w:t> descendant at any nesting depth. This works because deleted
text always uses the
DISTINCT <w:delText> tag rather than <w:t> -- so a plain <w:t> walk is
exactly the "all revisions accepted" reading, with no need to special-case
<w:ins>/<w:del> explicitly. This is a strict superset of the old behaviour:
a document with no tracked changes (e.g. the Facility Agreement) round-trips
identically either way, since every one of its <w:t> elements is already a
direct child of a <w:r> that is itself a direct child of <w:p>/<w:tc>.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _accepted_paragraph_text(p_element) -> str:
    """Accepted-revisions text for a single <w:p> element: every <w:t>
    descendant, regardless of how deeply it is nested inside <w:ins>
    wrappers. Deleted content uses <w:delText> (a different tag) and is
    therefore never collected."""
    return "".join(t.text or "" for t in p_element.iter(qn("w:t")))


def _accepted_cell_text(tc_element) -> str:
    """Accepted-revisions text for a <w:tc> table cell: its direct-child
    <w:p> paragraphs joined with "\\n", matching python-docx's own
    _Cell.text behaviour for multi-paragraph cells (walking every <w:t> in
    the whole cell subtree flat, with no separator, would silently glue
    adjacent paragraphs' words together)."""
    return "\n".join(
        _accepted_paragraph_text(p) for p in tc_element.findall(qn("w:p"))
    )


def extract(docx_path: Path) -> dict:
    """Walk a .docx's body in document order and return a JSON-serializable
    template: an ordered list of paragraph/table blocks. Placeholder markers
    ("[•]" or named "[Tokens]") and table structure are preserved verbatim
    -- no substitution happens here. Text reflects tracked changes as if
    every insertion were accepted and every deletion applied (see module
    docstring) -- never the raw, as-typed XML text.
    """
    document = docx.Document(str(docx_path))
    blocks: list[dict] = []
    for i, child in enumerate(document.element.body.iterchildren()):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            para = Paragraph(child, document)
            text = _accepted_paragraph_text(child)
            blocks.append(
                {
                    "type": "paragraph",
                    "index": i,
                    "style": para.style.name if para.style else None,
                    "text": text,
                    "placeholder_count": text.count("[•]"),
                }
            )
        elif tag == "tbl":
            table = Table(child, document)
            blocks.append(
                {
                    "type": "table",
                    "index": i,
                    "rows": [
                        [_accepted_cell_text(cell._tc) for cell in row.cells] for row in table.rows
                    ],
                }
            )
        # any other body-level tag (e.g. sectPr) carries no template content
        # and is intentionally skipped -- it is not a paragraph or a table.
    return {"slug": None, "source_file": docx_path.name, "blocks": blocks}


def main(argv: list[str]) -> None:
    if len(argv) < 2:
        raise SystemExit(
            "usage: extract_agreement_template.py <in.docx> <out.json> [slug]"
        )
    in_path, out_path = Path(argv[0]), Path(argv[1])
    slug = argv[2] if len(argv) > 2 else None
    result = extract(in_path)
    if slug:
        result["slug"] = slug
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"wrote {out_path} ({len(result['blocks'])} blocks)")


if __name__ == "__main__":
    main(sys.argv[1:])
