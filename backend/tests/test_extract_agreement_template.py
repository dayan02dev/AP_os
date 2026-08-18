"""Extractor tested against the real committed .docx (not a fixture).

If a revised source document changes shape (paragraph count, table layout,
placeholder count), these tests trip loudly — extraction must never silently
produce a truncated or misaligned template.
"""
from pathlib import Path

from scripts.extract_agreement_template import extract

SOURCE = (
    Path(__file__).resolve().parent.parent
    / "scripts"
    / "source_docs"
    / "facility_agreement_2026-08-06.docx"
)


def test_extracts_133_paragraphs_and_4_tables_in_order():
    result = extract(SOURCE)
    paras = [b for b in result["blocks"] if b["type"] == "paragraph"]
    tables = [b for b in result["blocks"] if b["type"] == "table"]
    assert len(paras) == 133
    assert len(tables) == 4
    # order preserved: the document's first block is the title paragraph
    assert result["blocks"][0]["text"].strip() == "FACILITY AGREEMENT"


def test_extracts_22_bullet_placeholders_total():
    result = extract(SOURCE)
    total = 0
    for b in result["blocks"]:
        if b["type"] == "paragraph":
            total += b["text"].count("[•]")
        else:
            for row in b["rows"]:
                for cell in row:
                    total += cell.count("[•]")
    assert total == 22


def test_collaborator_1_clause_is_extracted_verbatim():
    result = extract(SOURCE)
    para = next(b for b in result["blocks"] if b["type"] == "paragraph" and b["index"] == 4)
    assert para["placeholder_count"] == 4
    assert "Collaborator 1" in para["text"]


def test_facilities_schedule_table_has_six_placeholder_rows():
    result = extract(SOURCE)
    table = next(b for b in result["blocks"] if b["type"] == "table" and b["index"] == 126)
    placeholder_rows = [r for r in table["rows"] if "[•]" in r[3]]
    assert len(placeholder_rows) == 6
    assert table["rows"][1][1] == "Dedicated Seating"
    assert table["rows"][6][1] == "Administrative ID / Access Badge"


def test_a_revised_source_document_fails_loudly():
    """Guards against silent extraction drift: if the source changes shape,
    this test (run against the committed file) is the trip-wire — a
    paragraph or placeholder count that no longer matches is a signal the
    template JSON must be regenerated and every downstream index
    re-verified, not silently accepted."""
    result = extract(SOURCE)
    assert len(result["blocks"]) == 137  # 133 paragraphs + 4 tables


def test_paragraphs_only_traversal_would_lose_table_interleaving():
    """document.paragraphs (python-docx's flattened list) does not carry
    table position information — this test documents why extract() must
    walk document.element.body.iterchildren() instead."""
    import docx

    document = docx.Document(str(SOURCE))
    assert len(document.paragraphs) == 133  # same count...
    # ...but this list alone cannot tell you where the 4 tables sit relative
    # to these paragraphs in document order. extract()'s blocks list can.
    result = extract(SOURCE)
    table_positions = [i for i, b in enumerate(result["blocks"]) if b["type"] == "table"]
    assert table_positions == sorted(table_positions)
    assert len(table_positions) == 4


# ═══════════════════════════════════════════════════════════════════════════
# Collaboration Agreement — a REDLINED .docx with tracked changes. Its
# placeholders were rewritten mid-review: the old "[•]" markers sit inside
# <w:del> (deleted) runs and the new named tokens ("[Name of first
# founder]" etc.) sit inside <w:ins> (inserted) runs. python-docx's own
# Paragraph.text only walks DIRECT-CHILD <w:r> runs of <w:p> — it does not
# descend into <w:ins>/<w:del> wrapper elements — so it silently drops BOTH
# the deleted [•] markers AND the inserted named tokens, leaving orphaned
# fragments like "having PAN s/o/d/o resident of". extract() must instead
# walk every <w:t> descendant regardless of nesting depth: deleted text
# uses the distinct <w:delText> tag (never <w:t>), so a plain <w:t> walk
# naturally reconstructs the "revisions accepted" text without needing to
# special-case <w:ins>/<w:del> at all.
# ═══════════════════════════════════════════════════════════════════════════
COLLAB_SOURCE = (
    Path(__file__).resolve().parent.parent
    / "scripts"
    / "source_docs"
    / "collaboration_agreement_2026-08-15.docx"
)

_COLLAB_TOKENS = (
    "[Name of first founder]",
    "[PAN Number]",
    "[Father’s full name / Mother’s full name]",
    "[Address]",
    "[month]",
    "[date]",
    "[Date of agreement]",
    "[insert areas]",
)


def test_collaboration_agreement_extracts_173_paragraphs_and_2_tables_in_order():
    result = extract(COLLAB_SOURCE)
    paras = [b for b in result["blocks"] if b["type"] == "paragraph"]
    tables = [b for b in result["blocks"] if b["type"] == "table"]
    assert len(paras) == 173
    assert len(tables) == 2
    assert result["blocks"][0]["text"].strip() == "COLLABORATION AGREEMENT"


def test_collaboration_agreement_accepted_revisions_yields_no_orphaned_fragment():
    """The exact signature of getting tracked changes wrong: python-docx's
    naive .text on this document produces "having PAN s/o/d/o resident of"
    — the old [•] markers AND the new named tokens both vanish, because
    Paragraph.text doesn't descend into <w:ins>/<w:del>. Accepted-revisions
    extraction must never produce this fragment."""
    result = extract(COLLAB_SOURCE)
    full_text = "\n".join(b["text"] for b in result["blocks"] if b["type"] == "paragraph")
    assert "having PAN s/o/d/o resident of" not in full_text
    # and the real, accepted-revisions reading must be present instead
    assert (
        "[Name of first founder], having PAN [PAN Number], s/o/d/o "
        "[Father’s full name / Mother’s full name], resident of [Address]"
        in full_text
    )


def test_collaboration_agreement_has_16_named_placeholder_occurrences():
    result = extract(COLLAB_SOURCE)
    total = 0
    for b in result["blocks"]:
        if b["type"] == "paragraph":
            total += sum(b["text"].count(tok) for tok in _COLLAB_TOKENS)
        else:
            for row in b["rows"]:
                for cell in row:
                    total += sum(cell.count(tok) for tok in _COLLAB_TOKENS)
    assert total == 16


def test_collaboration_agreement_has_no_leftover_bullet_placeholder():
    """The old "[•]" convention was fully replaced by named tokens in the
    accepted redline — none should survive extraction."""
    result = extract(COLLAB_SOURCE)
    for b in result["blocks"]:
        if b["type"] == "paragraph":
            assert "[•]" not in b["text"]
        else:
            for row in b["rows"]:
                for cell in row:
                    assert "[•]" not in cell


def test_collaboration_agreement_signature_table_underscores_are_untouched():
    """The four blank signature lines are meant to stay blank — they are
    not placeholders to be filled by this pipeline."""
    result = extract(COLLAB_SOURCE)
    table = next(b for b in result["blocks"] if b["type"] == "table" and b["index"] == 161)
    underscore_cells = sum(
        1 for row in table["rows"] for cell in row if "___" in cell
    )
    assert underscore_cells == 4


def test_multi_paragraph_table_cells_keep_their_paragraph_breaks():
    """A table cell can hold more than one <w:p>. Walking every <w:t> in the
    whole cell subtree flat (no separator) would silently glue adjacent
    paragraphs' words together — extract() must join per-paragraph text
    with "\\n", matching python-docx's own _Cell.text behaviour, exactly
    as the un-redlined Schedule I premises table (index 122) requires."""
    result = extract(SOURCE)
    table = next(b for b in result["blocks"] if b["type"] == "table" and b["index"] == 122)
    assert table["rows"][1][1] == (
        "ARTPARK: ARTgarage facility (HMT, Jalahalli)\n"
        "Maximum of 500SqFT of facility will be allocated to the project"
    )


def test_extracting_facility_agreement_is_unaffected_by_tracked_change_support():
    """The Facility Agreement has no tracked changes at all — switching
    extract() to walk every <w:t> descendant (rather than relying on
    python-docx's shallow Paragraph.text) must be a no-op for it."""
    result = extract(SOURCE)
    paras = [b for b in result["blocks"] if b["type"] == "paragraph"]
    tables = [b for b in result["blocks"] if b["type"] == "table"]
    assert len(paras) == 133
    assert len(tables) == 4
    total = sum(b["text"].count("[•]") for b in result["blocks"] if b["type"] == "paragraph")
    total += sum(
        cell.count("[•]")
        for b in result["blocks"] if b["type"] == "table"
        for row in b["rows"] for cell in row
    )
    assert total == 22
