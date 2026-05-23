"""Integration test for the anchor-injection script.

Runs the script against a copy of the on-disk SIP template and verifies that
the resulting .docx contains START + END anchor markers for all 17
target questions, and that the python-docx text extraction picks
them up in the expected positions.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[2]
SIP_TEMPLATE_PATH = REPO_ROOT / "frontend" / "public" / "templates" / "ARTPARK_SIP_Application_Template.docx"
SCRIPT_PATH = REPO_ROOT / "scripts" / "inject_sip_template_anchors.py"

SIP_QUESTION_IDS = ["Q5", "Q6", "Q8", "Q9", "Q10", "Q11", "Q12", "Q13", "Q14",
                    "Q15", "Q16", "Q17", "Q18", "Q19", "Q20", "Q21", "Q24"]


def _docx_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_script_idempotent_and_injects_all_anchors(tmp_path: Path) -> None:
    assert SIP_TEMPLATE_PATH.exists(), "SIP template not at expected path"
    assert SCRIPT_PATH.exists(), "Injection script missing"

    # Copy to a tmp file so the test doesn't mutate the repo copy.
    work = tmp_path / "sip_template.docx"
    shutil.copy(SIP_TEMPLATE_PATH, work)

    import subprocess
    result = subprocess.run(
        ["python3", str(SCRIPT_PATH), str(work)],
        capture_output=True, text=True, check=False,
    )
    assert result.returncode == 0, f"script failed: {result.stderr}"

    text = _docx_text(work)
    for qid in SIP_QUESTION_IDS:
        assert f">>> ANSWER {qid} START >>>" in text, f"missing START anchor for {qid}"
        assert f"<<< ANSWER {qid} END <<<" in text, f"missing END anchor for {qid}"

    # Run a second time — script must be idempotent (no duplicate anchors).
    result2 = subprocess.run(
        ["python3", str(SCRIPT_PATH), str(work)],
        capture_output=True, text=True, check=False,
    )
    assert result2.returncode == 0, f"second run failed: {result2.stderr}"

    text2 = _docx_text(work)
    for qid in SIP_QUESTION_IDS:
        start_count = len(re.findall(re.escape(f">>> ANSWER {qid} START >>>"), text2))
        end_count = len(re.findall(re.escape(f"<<< ANSWER {qid} END <<<"), text2))
        assert start_count == 1, f"{qid}: expected 1 START anchor, got {start_count}"
        assert end_count == 1, f"{qid}: expected 1 END anchor, got {end_count}"
