"""mis_template_parser: parsing an uploaded MIS .docx into the catalog's own
shape, with a per-field record of what was matched, what was ambiguous and
what was ignored.

Two source fixtures, copied verbatim from the real ARTPARK templates
(docs/reference/mis-templates.md is the transcription; these .docx files are
ground truth per the build brief):
  - tests/fixtures/mis_monthly_template.docx
  - tests/fixtures/mis_quarterly_template.docx

Both fixtures are the BLANK templates ARTPARK ships — every content cell/
bullet is either empty, a bracketed placeholder ("[  ]"), or "e.g. ..."
instructional example text. Parsing them is the ground-truth test of
placeholder detection: nothing in a blank template may be imported as if it
were a founder's real answer. "Filled" variants are built in-memory by
mutating a fresh copy of the same real python-docx tree (same tables,
headings, styles — only the leaf text changes), so structural parsing is
still exercised against the real document shape, never a synthetic stand-in.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from app.services import mis_template_parser as parser
from app.services.mis_template_parser import MisParseError, parse_mis_document

FIXTURES = Path(__file__).parent / "fixtures"
MONTHLY_PATH = FIXTURES / "mis_monthly_template.docx"
QUARTERLY_PATH = FIXTURES / "mis_quarterly_template.docx"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _bytes(path: Path) -> bytes:
    return path.read_bytes()


def _doc(path: Path) -> Document:
    return Document(io.BytesIO(_bytes(path)))


def _save(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_cell(table, row: int, col: int, text: str) -> None:
    table.cell(row, col).text = text


def _set_para(doc: Document, index: int, text: str) -> None:
    p = doc.paragraphs[index]
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


# ── kind detection ──────────────────────────────────────────────────────

def test_detects_monthly_kind():
    result = parse_mis_document(file_bytes=_bytes(MONTHLY_PATH), mime=DOCX_MIME)
    assert result.kind == "monthly"


def test_detects_quarterly_kind():
    result = parse_mis_document(
        file_bytes=_bytes(QUARTERLY_PATH), mime=DOCX_MIME, fy_start_year=2026,
    )
    assert result.kind == "quarterly"


def test_rejects_unsupported_mime():
    with pytest.raises(MisParseError) as exc:
        parse_mis_document(file_bytes=b"not a docx", mime="application/pdf")
    assert exc.value.code == "unsupported_mime"


def test_rejects_empty_file():
    with pytest.raises(MisParseError) as exc:
        parse_mis_document(file_bytes=b"", mime=DOCX_MIME)
    assert exc.value.code == "empty_document"


def test_quarterly_without_fy_start_year_raises():
    with pytest.raises(MisParseError) as exc:
        parse_mis_document(file_bytes=_bytes(QUARTERLY_PATH), mime=DOCX_MIME)
    assert exc.value.code == "fy_start_year_required"


# ── monthly: blank template -> everything reads as blank, nothing invented ──

def test_blank_monthly_template_has_no_ambiguous_or_ignored_metrics():
    result = parse_mis_document(file_bytes=_bytes(MONTHLY_PATH), mime=DOCX_MIME)
    # Every one of the 13 catalog metrics is located (matched) with a real
    # None for both target and actual -- the template's "[  ]"/"e.g. ..."
    # cells are placeholders, not data.
    assert len(result.metrics) == 13
    for m in result.metrics:
        assert m["actual"] is None
        # product_metric_1/2's Target column is itself "e.g. accuracy, uptime"
        # in the source -- instructional, not a value.
        assert m["target"] is None
    assert not result.ambiguous
    assert not result.ignored


def test_blank_monthly_template_milestones_table_yields_no_rows():
    result = parse_mis_document(file_bytes=_bytes(MONTHLY_PATH), mime=DOCX_MIME)
    # All 6 rows are "[milestone description]"-style placeholders -- a blank
    # row is not a milestone.
    assert result.entries.get("milestones", []) == []


def test_blank_monthly_template_narrative_is_all_blank():
    result = parse_mis_document(file_bytes=_bytes(MONTHLY_PATH), mime=DOCX_MIME)
    assert result.narrative == {}


# ── monthly: filled metrics table ────────────────────────────────────────

def test_filled_monthly_metrics_table_parses_real_numbers():
    doc = _doc(MONTHLY_PATH)
    tbl = doc.tables[0]
    _set_cell(tbl, 2, 1, "10")   # revenue_month target
    _set_cell(tbl, 2, 2, "12.5")  # revenue_month actual
    _set_cell(tbl, 2, 4, "Ahead of plan")
    _set_cell(tbl, 10, 2, "7")   # TRL row -- founder must never be able to set this
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    by_key = {m["metric_key"]: m for m in result.metrics}
    assert by_key["revenue_month"]["target"] == 10.0
    assert by_key["revenue_month"]["actual"] == 12.5
    assert by_key["revenue_month"]["commentary"] == "Ahead of plan"
    # trl_level is server-computed; the parser must never surface a founder-
    # supplied value for it even if the cell was edited.
    assert "trl_level" not in by_key or by_key["trl_level"]["actual"] is None
    assert any(n.path == "metrics.trl_level.actual" for n in result.ignored)


def test_filled_monthly_metrics_table_flags_unparseable_number_as_ambiguous():
    doc = _doc(MONTHLY_PATH)
    tbl = doc.tables[0]
    _set_cell(tbl, 2, 2, "twelve point five")  # revenue_month actual, garbled
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    by_key = {m["metric_key"]: m for m in result.metrics}
    assert by_key["revenue_month"]["actual"] is None
    assert any(
        n.path == "metrics.revenue_month.actual" and n.reason == "unparseable_number"
        for n in result.ambiguous
    )


def test_metrics_row_not_in_catalog_is_ignored_not_coerced():
    doc = _doc(MONTHLY_PATH)
    tbl = doc.tables[0]
    _set_cell(tbl, 2, 0, "Some Business-Specific KPI Founders Added")
    _set_cell(tbl, 2, 2, "99")
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    keys = {m["metric_key"] for m in result.metrics}
    assert "revenue_month" not in keys  # row 2 no longer matches this key
    assert any(n.reason == "unmatched_row" for n in result.ignored)


# ── monthly: milestones table ────────────────────────────────────────────

def test_filled_milestones_row_parses():
    doc = _doc(MONTHLY_PATH)
    tbl = doc.tables[1]
    _set_cell(tbl, 1, 0, "Tape out v0.3")
    _set_cell(tbl, 1, 1, "Priya")
    _set_cell(tbl, 1, 2, "On Track")
    _set_cell(tbl, 1, 3, "Slipped two weeks")
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    rows = result.entries["milestones"]
    assert len(rows) == 1
    assert rows[0] == {
        "milestone": "Tape out v0.3", "owner": "Priya",
        "status": "On Track", "notes": "Slipped two weeks",
    }


def test_milestone_invalid_status_is_ambiguous():
    doc = _doc(MONTHLY_PATH)
    tbl = doc.tables[1]
    _set_cell(tbl, 1, 0, "Tape out v0.3")
    _set_cell(tbl, 1, 2, "Completed")  # not one of the 4 catalog choices
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    rows = result.entries["milestones"]
    assert len(rows) == 1
    assert "status" not in rows[0]
    assert any(
        n.path.startswith("entries.milestones[0].status") and n.reason == "invalid_value"
        for n in result.ambiguous
    )


# ── monthly: narrative label matching ────────────────────────────────────

def test_filled_exec_summary_narrative():
    doc = _doc(MONTHLY_PATH)
    _set_para(doc, 6, "Headline win: shipped v2 to three hospitals.")
    _set_para(doc, 9, "Cash: 4.2 Cr in bank, 9 months runway.")
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    assert result.narrative["exec.headline_win"] == "shipped v2 to three hospitals."
    assert result.narrative["exec.cash"] == "4.2 Cr in bank, 9 months runway."
    assert "exec.biggest_concern" not in result.narrative


def test_filled_team_hiring_narrative_label_has_extra_words_in_doc():
    doc = _doc(MONTHLY_PATH)
    # Doc label is "Open roles (urgent first):" -- catalog prompt is just
    # "Open roles". The parser must match on the real document label, not
    # the catalog's own prompt string.
    _set_para(doc, 39, "Open roles (urgent first): 2x firmware engineers.")
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    assert result.narrative["team.open_roles"] == "2x firmware engineers."


# ── monthly: risks (prose entries with embedded Impact:/Mitigation: anchors) ──

def test_filled_risk_bullet_decomposes_on_impact_mitigation_anchors():
    doc = _doc(MONTHLY_PATH)
    _set_para(
        doc, 31,
        "Risk 1 (red / blocked): ASIC yield at 61%. "
        "Impact: cost per unit rose to $28. "
        "Mitigation: DFT rework in progress.",
    )
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    rows = result.entries["risks"]
    assert len(rows) == 1
    assert rows[0]["severity"] == "red"
    assert rows[0]["what_happened"] == "ASIC yield at 61%."
    assert rows[0]["impact"] == "cost per unit rose to $28."
    assert rows[0]["mitigation"] == "DFT rework in progress."


# ── monthly: asks (numbered list, category has no textual anchor) ────────

def test_filled_ask_bullet_parses_priority_and_text_leaves_category_ambiguous():
    doc = _doc(MONTHLY_PATH)
    _set_para(doc, 53, "1. Intro to VP Ops at Fortis Healthcare.")
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME)
    rows = result.entries["asks"]
    assert len(rows) == 1
    assert rows[0]["priority"] == 1
    assert rows[0]["ask"] == "Intro to VP Ops at Fortis Healthcare."
    assert "category" not in rows[0]
    assert any(
        n.path.startswith("entries.asks[0].category") for n in result.ambiguous
    )


# ── quarterly: financials tables + self-contained needs_gap cross-check ──

def test_blank_quarterly_financials_and_headcount_are_all_none():
    result = parse_mis_document(
        file_bytes=_bytes(QUARTERLY_PATH), mime=DOCX_MIME, fy_start_year=2026,
    )
    assert len(result.financials) > 0
    assert all(row["amount"] is None for row in result.financials)
    assert len(result.headcount) == 4  # Total row excluded
    assert all(row["current_count"] is None for row in result.headcount)
    assert not result.ambiguous


def test_filled_annual_revenue_table_maps_doc_fy_labels_to_catalog_buckets():
    doc = _doc(QUARTERLY_PATH)
    tbl = doc.tables[0]
    _set_cell(tbl, 1, 1, "120")   # booked, first historical bucket
    _set_cell(tbl, 1, 5, "80")    # booked, YTD bucket
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    from app.services import mis_catalog as cat
    buckets = cat.annual_revenue_buckets(2026)
    by_bucket = {
        r["bucket"]: r["amount"] for r in result.financials
        if r["series"] == "annual_revenue_booked"
    }
    assert by_bucket[buckets[0]] == 120.0
    assert by_bucket[buckets[4]] == 80.0  # "... YTD" bucket


def test_filled_needs_table_gap_mismatch_is_cross_checked_never_persisted():
    doc = _doc(QUARTERLY_PATH)
    tbl = doc.tables[1]
    _set_cell(tbl, 1, 1, "100")  # total, Q1
    _set_cell(tbl, 2, 1, "40")   # confirmed, Q1
    _set_cell(tbl, 3, 1, "30")   # projected, Q1
    _set_cell(tbl, 4, 1, "999")  # gap, Q1 -- deliberately wrong (true gap is 30)
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    series = {r["series"] for r in result.financials}
    assert "needs_gap" not in series  # never persisted, even though present in the doc
    mismatches = [c for c in result.cross_check_mismatches if c.path.startswith("financials.needs_gap")]
    assert mismatches, "expected the doc's Gap cell to be flagged against the computed value"
    assert mismatches[0].doc_value == 999.0
    assert mismatches[0].computed_value == 30.0


def test_filled_needs_table_gap_matching_computed_value_is_not_flagged():
    doc = _doc(QUARTERLY_PATH)
    tbl = doc.tables[1]
    _set_cell(tbl, 1, 1, "100")
    _set_cell(tbl, 2, 1, "40")
    _set_cell(tbl, 3, 1, "30")
    _set_cell(tbl, 4, 1, "30")  # correct gap
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    mismatches = [c for c in result.cross_check_mismatches if c.path.startswith("financials.needs_gap")]
    assert mismatches == []


# ── quarterly: headcount table, Total row + Net Change never stored ─────

def test_filled_headcount_table_excludes_total_row_and_net_change_column():
    doc = _doc(QUARTERLY_PATH)
    tbl = doc.tables[2]
    _set_cell(tbl, 1, 1, "12")   # artpark_associated current_count
    _set_cell(tbl, 1, 2, "1")    # exited
    _set_cell(tbl, 1, 3, "3")    # Net Change column -- must never be read into a stored field
    _set_cell(tbl, 5, 1, "999")  # Total row current_count -- must never become a category row
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    assert len(result.headcount) == 4
    by_cat = {r["category"]: r for r in result.headcount}
    assert "total" not in by_cat
    assert by_cat["artpark_associated"]["current_count"] == 12
    assert by_cat["artpark_associated"]["exited"] == 1
    assert "net_change" not in by_cat["artpark_associated"]
    assert not any(row["current_count"] == 999 for row in result.headcount)


def test_headcount_net_change_cross_check_against_supplied_previous_period():
    doc = _doc(QUARTERLY_PATH)
    tbl = doc.tables[2]
    _set_cell(tbl, 1, 1, "12")  # current_count
    _set_cell(tbl, 1, 3, "3")   # doc's own (wrong) Net Change
    result = parse_mis_document(
        file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026,
        prev_headcount_current={"artpark_associated": 15},  # true net change = 12 - 15 = -3
    )
    mismatches = [
        c for c in result.cross_check_mismatches
        if c.path == "headcount.net_change.artpark_associated"
    ]
    assert mismatches
    assert mismatches[0].doc_value == 3.0
    assert mismatches[0].computed_value == -3.0


# ── monthly: vs_last cross-check against a supplied previous actual ──────

def test_vs_last_cross_check_against_supplied_previous_actual():
    doc = _doc(MONTHLY_PATH)
    tbl = doc.tables[0]
    _set_cell(tbl, 2, 2, "12")  # actual
    _set_cell(tbl, 2, 3, "5")   # vs Last Mo doc value -- wrong, true is 12-8=4
    result = parse_mis_document(
        file_bytes=_save(doc), mime=DOCX_MIME,
        prev_actual_by_metric_key={"revenue_month": 8},
    )
    mismatches = [c for c in result.cross_check_mismatches if c.path == "metrics.vs_last.revenue_month"]
    assert mismatches
    assert mismatches[0].doc_value == 5.0
    assert mismatches[0].computed_value == 4.0


# ── quarterly: planned-vs-actual (9.1) decomposes on its own embedded labels ──

def test_quarterly_planned_vs_actual_decomposes_on_embedded_labels():
    doc = _doc(QUARTERLY_PATH)
    _set_para(
        doc, 88,
        "Planned: Silicon v0.2 tape-out at 75% yield. Actual: yield 61%. "
        "Reason: DFT gaps. Corrective action: yield-improvement pass.",
    )
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    rows = result.entries["planned_vs_actual"]
    assert len(rows) == 1
    assert rows[0]["planned"] == "Silicon v0.2 tape-out at 75% yield."
    assert rows[0]["achieved"] == "yield 61%."
    assert rows[0]["reason"] == "DFT gaps."
    assert rows[0]["corrective_action"] == "yield-improvement pass."
    assert "outcome" not in rows[0]  # no textual anchor for outcome -- never guessed


# ── quarterly: bucketed prose entries are never decomposed, only flagged ──

def test_quarterly_ip_register_prose_is_flagged_not_fabricated():
    doc = _doc(QUARTERLY_PATH)
    _set_para(
        doc, 19,
        "Patent — Novel Doppler algorithm for portable ultrasound "
        "(Healthcare, National, IN2026XXXXXX, filed Feb-2026).",
    )
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    # Never invented: no structured ip_assets row is fabricated from prose.
    assert result.entries.get("ip_assets", []) == []
    notes = [n for n in result.ambiguous if n.path.startswith("entries.ip_assets")]
    assert notes
    assert notes[0].raw is not None and "Doppler" in notes[0].raw


# ── mutation-relevant: metrics/vs_last & headcount/net_change sign safety ──

def test_needs_gap_cross_check_uses_subtraction_not_addition():
    """Locks the exact arithmetic (total - confirmed - projected), matching
    mis_catalog's documented rule, so a future edit cannot silently flip the
    sign the way the headcount net_change bug once did."""
    doc = _doc(QUARTERLY_PATH)
    tbl = doc.tables[1]
    _set_cell(tbl, 1, 1, "50")
    _set_cell(tbl, 2, 1, "50")
    _set_cell(tbl, 3, 1, "50")
    _set_cell(tbl, 4, 1, "-50")  # correct gap: 50 - 50 - 50 = -50
    result = parse_mis_document(file_bytes=_save(doc), mime=DOCX_MIME, fy_start_year=2026)
    mismatches = [c for c in result.cross_check_mismatches if c.path.startswith("financials.needs_gap")]
    assert mismatches == []
