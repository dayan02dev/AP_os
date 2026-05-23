#!/usr/bin/env python3
"""Generate the six SIP template test fixtures from a small data spec.

Output files land in backend/tests/fixtures/.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
FIXTURE_DIR = REPO_ROOT / "backend" / "tests" / "fixtures"
SOURCE_TEMPLATE = REPO_ROOT / "frontend" / "public" / "templates" / "ARTPARK_SIP_Application_Template.docx"

ALL_QIDS = ["Q5", "Q6", "Q8", "Q9", "Q10", "Q11", "Q12", "Q13", "Q14",
            "Q15", "Q16", "Q17", "Q18", "Q19", "Q20", "Q21", "Q24"]

REQUIRED_QIDS = ["Q5", "Q6", "Q8", "Q10", "Q11", "Q12", "Q13", "Q15",
                 "Q16", "Q17", "Q18", "Q19"]

COMPLETE_ANSWERS = {
    "Q5":  "Yes — Pvt Ltd, registered in India",
    "Q6":  "TRL 4 — lab-validated prototype",
    "Q8":  "No",
    "Q9":  "",
    "Q10": "Referral from friend/colleague",
    "Q11": "We tackle the problem of inadequate post-harvest cold chain in tier-3 towns.",
    "Q12": "A modular, solar-powered cold storage unit deployable in 48 hours.",
    "Q13": "Phase-change thermal storage built from a patented composite.",
    "Q14": "Most experts overweight capex; opex dominates 5-year TCO.",
    "Q15": "Active pilots (paid or unpaid) with design partners",
    "Q16": "Two pilots with FPOs in Maharashtra, one paid LOI in Karnataka.",
    "Q17": "Material fatigue under 45C ambient — solved with composite v2.",
    "Q18": "Three deployments by Q3; one paying customer by Q4; opex reduced 40% YoY.",
    "Q19": "Specialized sensor calibration lab; access to robotics testbed.",
    "Q20": "Initial prototype failed in monsoon humidity; pivoted to sealed composite.",
    "Q21": "Strict hardware-in-the-loop sim suite; weekly drift audits.",
    "Q24": "https://www.loom.com/share/abc123",
}


def _write_answers(path: Path, answers: dict[str, str]) -> None:
    shutil.copy(SOURCE_TEMPLATE, path)
    doc = Document(str(path))

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paras = cell.paragraphs
                for i, p in enumerate(paras):
                    txt = p.text.strip()
                    if not txt.startswith(">>> ANSWER "):
                        continue
                    qid = txt.split(" ")[2]
                    if qid not in answers:
                        continue
                    end_idx = None
                    for j in range(i + 1, len(paras)):
                        if paras[j].text.strip().startswith("<<< ANSWER "):
                            end_idx = j
                            break
                    if end_idx is None:
                        continue
                    answer = answers[qid]
                    for j in range(i + 1, end_idx):
                        paras[j].text = ""
                    if i + 1 < end_idx:
                        paras[i + 1].text = answer

    doc.save(str(path))


def build_anchored_complete(path: Path) -> None:
    _write_answers(path, COMPLETE_ANSWERS)


def build_anchored_partial(path: Path) -> None:
    partial = {k: v for k, v in COMPLETE_ANSWERS.items() if k in REQUIRED_QIDS}
    _write_answers(path, partial)


def build_empty(path: Path) -> None:
    shutil.copy(SOURCE_TEMPLATE, path)
    doc = Document(str(path))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip().startswith((">>> ANSWER", "<<< ANSWER")):
                        p.text = ""
    doc.save(str(path))


def build_anchors_stripped(path: Path) -> None:
    build_anchored_complete(path)
    doc = Document(str(path))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip().startswith((">>> ANSWER", "<<< ANSWER")):
                        p.text = ""
    doc.save(str(path))


def main() -> int:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    build_anchored_complete(FIXTURE_DIR / "sip_template_anchored_complete.docx")
    build_anchored_partial(FIXTURE_DIR / "sip_template_anchored_partial.docx")
    build_empty(FIXTURE_DIR / "sip_template_empty.docx")
    build_anchors_stripped(FIXTURE_DIR / "sip_template_anchors_stripped.docx")

    shutil.copy(
        REPO_ROOT / "frontend" / "public" / "templates" / "ARTPARK_TIR_Application_Template.docx",
        FIXTURE_DIR / "sip_template_tir_uploaded.docx",
    )

    print(f"wrote 5 fixtures to {FIXTURE_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
