#!/usr/bin/env python3
"""Inject anchor markers into the SIP application template .docx.

The SIP template ships without `>>> ANSWER QN START >>>` / `<<< ANSWER QN
END <<<` markers, but the SIP parser (services/sip_template_parser.py)
depends on them. This script walks the document, finds each question
heading, and inserts marker paragraphs inside the corresponding answer
cell so the parser can deterministically slice each answer.

Usage:
    python3 scripts/inject_sip_template_anchors.py <path/to/template.docx>

Idempotent: a second run is a no-op (markers already present are detected
and skipped).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

TARGET_QIDS = ["Q5", "Q6", "Q8", "Q9", "Q10", "Q11", "Q12", "Q13", "Q14",
               "Q15", "Q16", "Q17", "Q18", "Q19", "Q20", "Q21", "Q24"]

HEADING_RE = re.compile(
    r"^\s*(Q\d+)\s*[··]\s*(REQUIRED|OPTIONAL)\b",
    re.IGNORECASE,
)


def _heading_qid(text: str) -> str | None:
    m = HEADING_RE.match(text)
    if not m:
        return None
    qid = m.group(1).upper()
    return qid if qid in TARGET_QIDS else None


def _cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs)


def _insert_marker_at_start(cell: _Cell, marker: str) -> None:
    """Insert a paragraph at the top of the cell."""
    cell.paragraphs[0].insert_paragraph_before(marker)


def _insert_marker_at_end(cell: _Cell, marker: str) -> None:
    """Append a paragraph at the bottom of the cell."""
    cell.add_paragraph(marker)


def _find_following_table(body_children: list, after_paragraph_element):
    """Return the next w:tbl element after the given w:p element."""
    seen_p = False
    for el in body_children:
        if el is after_paragraph_element:
            seen_p = True
            continue
        if not seen_p:
            continue
        if el.tag == qn("w:tbl"):
            return el
        if el.tag == qn("w:p"):
            text = "".join(t.text or "" for t in el.iter(qn("w:t")))
            if HEADING_RE.match(text):
                return None
    return None


def inject(path: Path) -> int:
    """Inject anchors in-place. Returns the number of question pairs added."""
    doc = Document(str(path))
    body = doc.element.body
    body_children = list(body)

    para_by_el = {p._p: p for p in doc.paragraphs}
    table_by_el = {t._tbl: t for t in doc.tables}

    pairs_added = 0
    for i, el in enumerate(body_children):
        if el.tag != qn("w:p"):
            continue
        p = para_by_el.get(el)
        if p is None:
            continue
        qid = _heading_qid(p.text)
        if not qid:
            continue

        next_tbl = _find_following_table(body_children, el)
        if next_tbl is None:
            print(f"[warn] no answer table found after heading {qid}", file=sys.stderr)
            continue

        tbl = table_by_el.get(next_tbl)
        if tbl is None:
            print(f"[warn] table found but not wrapped for {qid}", file=sys.stderr)
            continue

        cell = tbl.rows[0].cells[0]

        existing = _cell_text(cell)
        start_marker = f">>> ANSWER {qid} START >>>"
        end_marker = f"<<< ANSWER {qid} END <<<"

        if start_marker in existing and end_marker in existing:
            continue

        if start_marker not in existing:
            _insert_marker_at_start(cell, start_marker)
        if end_marker not in existing:
            _insert_marker_at_end(cell, end_marker)
        pairs_added += 1

    doc.save(str(path))
    return pairs_added


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        print("usage: inject_sip_template_anchors.py <path/to/template.docx>", file=sys.stderr)
        return 2
    path = Path(argv[1])
    if not path.exists():
        print(f"file not found: {path}", file=sys.stderr)
        return 3
    n = inject(path)
    print(f"injected/verified {n} question pair(s) in {path}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
